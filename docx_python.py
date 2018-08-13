import ntpath
import zipfile
from mailmerge import MailMerge
import json
import subprocess
import os
from docx.shared import Inches
from docx import Document
import shutil
import xml.etree.ElementTree as ET


class DOCXPython():

    def __init__(self, project_path='/Users/minhchau/Downloads/DocxPython', word_template_name='word_template', convert_to_pdf = True,
                 multiple_files = True):
        self.project_path = project_path
        self.convert_to_pdf = convert_to_pdf
        self.multiple_files = multiple_files

        self.input_data = os.path.join(self.project_path, 'input_data')
        self.data_mailingMerge = os.path.join(self.input_data, 'data_mailingMerge.json')
        self.data_signature = os.path.join(self.input_data, 'data_signature')
        self.data_signature_json = os.path.join(self.data_signature, 'data_signature.json')
        self.data_header = os.path.join(self.input_data, 'data_header')
        self.data_footer = os.path.join(self.input_data, 'data_footer')
        self.data_signature_image = os.path.join(self.data_signature, 'signature.png')

        self.word_template_name = word_template_name
        self.word_template = os.path.join(self.input_data, self.word_template_name)

        self.output = os.path.join(self.project_path, 'output')
        self.output_Word = os.path.join(self.output, 'WORD')
        self.output_PDF = os.path.join(self.output, 'PDF')

        self.name_key = '«name»'
        self.signature_key = '«signature»'
        self.signature_width = 1.25  # inches
        self.signature_height = None

        self.word_extension = '.docx'
        self.fileName_prefix = self.word_template_name.replace(self.word_extension, '') + '_'
        self.word_extension = '.docx'

    def replace_header_and_footer(self):

        # helper function
        def path_leaf(path):
            head, tail = ntpath.split(path)
            return tail or ntpath.basename(head)

        TEMP_ZIP = os.path.join(self.input_data, 'template.docx.zip')
        TEMP_FOLDER = os.path.join(self.input_data, 'template')

        # remove old zip file or folder template
        if os.path.exists(TEMP_ZIP):
            os.remove(TEMP_ZIP)

        if os.path.exists(TEMP_FOLDER):
            shutil.rmtree(TEMP_FOLDER)

        # reformat template.docx's extension
        os.rename(self.word_template, TEMP_ZIP)

        # unzip file zip to specific folder
        with zipfile.ZipFile(TEMP_ZIP, 'r') as z:
            z.extractall(TEMP_FOLDER)

        files_in_rels = os.listdir(os.path.join(TEMP_FOLDER, 'word', '_rels'))

        # HEADER:
        for file_name in files_in_rels:
            if 'header' in file_name:
                header_file = file_name
                break

        header_xml = os.path.join(TEMP_FOLDER, 'word', '_rels', header_file)
        # todo: Make 'header2.xml.rels' dynamic (can be different when using other templates)

        tree = ET.parse(header_xml)
        root = tree.getroot()

        header_name_list = os.listdir(self.data_header)
        try:
            header_name_list.remove('.DS_Store')
        except:
            pass
        finally:
            header_name = header_name_list[0]

        child = root.getchildren()[0].attrib
        child['Target'] = os.path.join('media', header_name)
        tree.write(header_xml)

        # FOOTER:
        for file_name in files_in_rels:
            if 'footer' in file_name:
                footer_file = file_name
                break

        footer_xml = os.path.join(TEMP_FOLDER, 'word', '_rels', footer_file)
        tree_footer = ET.parse(footer_xml)
        root_footer = tree_footer.getroot()

        footer_name_list = os.listdir(self.data_footer)
        try:
            footer_name_list.remove('.DS_Store')
        except:
            pass
        finally:
            footer_name = footer_name_list[0]

        child = root_footer.getchildren()[0].attrib
        child['Target'] = os.path.join('media', footer_name)
        tree_footer.write(footer_xml)

        media_folder = os.path.join(TEMP_FOLDER, 'word', 'media')
        shutil.copy2(os.path.join(self.data_header, header_name), media_folder)
        shutil.copy2(os.path.join(self.data_footer, footer_name), media_folder)

        # todo: delete former footer, header

        # zip temp folder to zip file
        os.remove(TEMP_ZIP)
        shutil.make_archive(TEMP_ZIP.replace('.zip', ''), 'zip', TEMP_FOLDER)

        # rename zip file to docx
        os.rename(TEMP_ZIP, self.word_template)
        shutil.rmtree(TEMP_FOLDER)

        print('replace_header_and_footer is done')

    def replace_signature(self):
        with open(self.data_signature_json) as f:
            data = json.load(f)

        NAME = data[self.name_key]
        SIGNATURE = os.path.join(self.data_signature, data[self.signature_key])

        word_templates = os.listdir(self.input_data)

        document = Document(self.word_template)

        sections = document.paragraphs
        len_sections = len(sections)

        signature_replaced = False
        name_replaced = False

        for i in range(len_sections):
            p = sections[len_sections - i - 1]

            if self.signature_key in p.text:
                new_text = p.text.replace(self.signature_key, '')
                p.clear()
                r = p.add_run()
                r.add_text(new_text)
                r.add_picture(SIGNATURE, width=Inches(self.signature_width))
                signature_replaced = True

            if self.name_key in p.text:
                p.text = p.text.replace(self.name_key, NAME)
                name_replaced = True

            if signature_replaced == True and name_replaced == True:
                break

        document.save(self.word_template)
        print('replace_signature is done !')

    def replace_mailing_merge_and_convert_to_pdf(self):

        with open(self.data_mailingMerge) as f:
            data = json.load(f)['data']
        len_data = len(data)

        if not os.path.exists(self.output_Word):
            os.makedirs(self.output_Word)

        if not os.path.exists(self.output_PDF):
            os.makedirs(self.output_PDF)

        if self.multiple_files:
            for index, input in enumerate(data):
                print('creating {}/{} ...'.format(index, len_data))
                document = MailMerge(self.word_template)
                document.merge(**input)

                word_output = os.path.join(self.output_Word, self.fileName_prefix + str(index + 1) + self.word_extension)
                document.write(word_output)

                # convert to PDF
                if self.convert_to_pdf:
                    subprocess.check_output(
                        ['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--convert-to', 'pdf', '--outdir',
                         self.output_PDF,
                         word_output])
        else:
            print("creating one file containing {} mails...".format(len_data))
            document = MailMerge(self.word_template)
            document.merge_pages(data)

            word_output = os.path.join(self.output_Word, self.fileName_prefix + self.word_extension)
            document.write(word_output)

            # convert to PDF
            if self.convert_to_pdf:
                subprocess.check_output(
                    ['/Applications/LibreOffice.app/Contents/MacOS/soffice', '--convert-to', 'pdf', '--outdir',
                     self.output_PDF,
                     word_output])


        print('replace_mailingMerge_and_convertToPDF is done !')

    def print_finish(self):
        print('===========================')

    def get_new_word_template(self, data_source = None, delete_output_folder = True):
        if data_source is not None:
            src = os.path.join(self.project_path, 'other_templates', data_source)
            dst = os.path.join(self.input_data, data_source)
            shutil.copy2(src, dst)

        try:
            if delete_output_folder:
                shutil.rmtree(self.output)
        except:
            pass
        print('-- get a new word template -- ')

    @classmethod
    def clear_all_data(cls, project_path):

        output = os.path.join(project_path, 'output')
        input_data = os.path.join(project_path, 'input_data')

        do_not_delete = ['data_header', 'data_mailingMerge.json', 'data_signature', 'word_template.docx', 'data_footer']
        files_in_input_data_folder = os.listdir(input_data)
        try:
            for file_name in files_in_input_data_folder:
                if file_name in do_not_delete:
                    pass
                else:
                    print("os.path.join(input_data, file_name)" + str(os.path.join(input_data, file_name)))
                    os.remove(os.path.join(input_data, file_name))
            shutil.rmtree(output)

        except:
            pass

        print('-- clear all data -- ')

